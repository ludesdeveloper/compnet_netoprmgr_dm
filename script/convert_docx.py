from docx import Document
import sqlite3
import re

class convert_docx:
    @staticmethod
    def convert_docx():
        print('')
        print('Processing Document')

        #open db connection
        db = sqlite3.connect('pmdb')
        cursor = db.cursor()
        #using document docx module
        document = Document()
        
        #SOFTWARE SUMMARY TABLE
        #sql query
        cursor.execute('''SELECT version, COUNT(*) FROM swsumtable GROUP BY version''')
        records = cursor.fetchall()
        cursor.execute('''SELECT COUNT(version) FROM swsumtable''')
        total = cursor.fetchall()
        total=(str(total))
        total=re.sub("\D", "", total)
        total=int(total)
        #add to document
        p = document.add_paragraph('')
        p.add_run('SOFTWARE TABLE SUMMARY').bold = True
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Version'
        hdr_cells[1].text = 'Total'
        hdr_cells[2].text = 'Percentage'        
        for row in records:
            row_cells = table.add_row().cells
            row_cells[0].text = (row[0])
            row_cells[1].text = str(row[1]) 
            percentage = (str((row[1]/total)*100))
            percentage = re.findall('\d+[.]\d',percentage)
            percentage = percentage[0]
            row_cells[2].text = (percentage+'%')
         
        #SOFTWARE TABLE
        #sql query
        cursor.execute('''SELECT devicename, model, iosversion, uptime, confreg FROM swtable''')
        records = cursor.fetchall()
        #add to document
        p = document.add_paragraph('')
        p.add_run('SOFTWARE TABLE').bold = True
        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Device Name'
        hdr_cells[1].text = 'Model'
        hdr_cells[2].text = 'IOS Version'
        hdr_cells[3].text = 'Uptime'
        hdr_cells[4].text = 'Config Register'
        for row in records:
            row_cells = table.add_row().cells
            row_cells[0].text = (row[0])
            row_cells[1].text = (row[1])
            row_cells[2].text = (row[2])
            row_cells[3].text = (row[3])
            row_cells[4].text = (row[4])           
        
        #HARDWARE SUMMARY TABLE
        #sql query
        cursor.execute('''SELECT model, COUNT(*) FROM hwsumtable GROUP BY model''')
        records = cursor.fetchall()
        cursor.execute('''SELECT COUNT(model) FROM hwsumtable''')
        total = cursor.fetchall()
        total=(str(total))
        total=re.sub("\D", "", total)
        total=int(total)
        #add to document
        p = document.add_paragraph('')
        p.add_run('HARDWARE TABLE SUMMARY').bold = True
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Model'
        hdr_cells[1].text = 'Total'
        hdr_cells[2].text = 'Percentage'        
        for row in records:
            row_cells = table.add_row().cells
            row_cells[0].text = (row[0])
            row_cells[1].text = str(row[1])
            percentage = (str((row[1]/total)*100))
            percentage = re.findall('\d+[.]\d',percentage)
            percentage = percentage[0]
            row_cells[2].text = (percentage+'%')
         
        #HARDWARE CARD TABLE
        #sql query
        cursor.execute('''SELECT devicename, model, card, sn, hwdscr FROM hwcardtable''')
        records = cursor.fetchall()
        p = document.add_paragraph('')
        p.add_run('HARDWARE CARD TABLE').bold = True

        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Device Name'
        hdr_cells[1].text = 'Model'
        hdr_cells[2].text = 'Card'
        hdr_cells[3].text = 'Description'
        hdr_cells[4].text = 'Serial Number'
        iteration_check = False
        row_check = 'ludesdeveloper'
        count_row = 1
        for row in records:
            row_cells = table.add_row().cells
            if row_check not in row[0]:
                row_check = row[0]
                iteration_check = False
                start_row = table.cell(count_row, 0)
            else:
                pass
            if row_check == row[0] and iteration_check == False:
                row_cells[0].text = (row[0])
                iteration_check = True
            else:
                row_cells[0].text = ('')
                end_row = table.cell(count_row, 0)
                merge_row = start_row.merge(end_row)
            row_cells[1].text = (row[1])
            row_cells[2].text = (row[2])
            row_cells[3].text = (row[4])
            row_cells[4].text = (row[3])
            count_row+=1
         
        #CPU SUMMARY TABLE
        #sql query
        cursor.execute('''SELECT devicename, model, total, process, interrupt, topcpu, status FROM cpusumtable''')
        records = cursor.fetchall()
        p = document.add_paragraph('')
        p.add_run('CPU SUMMARY TABLE').bold = True

        table = document.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Device Name'
        hdr_cells[1].text = 'Model'
        hdr_cells[2].text = 'Total'
        hdr_cells[3].text = 'Process'
        hdr_cells[4].text = 'Interrupt'
        hdr_cells[5].text = 'Top Process'
        hdr_cells[6].text = 'Status'

        for row in records:

            row_cells = table.add_row().cells
            row_cells[0].text = (row[0])
            row_cells[1].text = (row[1])
            row_cells[2].text = (row[2]+"%")
            row_cells[3].text = (row[3]+"%")
            row_cells[4].text = (row[4]+"%")
            row_cells[5].text = (row[5])
            row_cells[6].text = (row[6])
        
        #MEMORY SUMMARY TABLE
        #sql query
        cursor.execute('''SELECT devicename, model, utils, topproc, status FROM memsumtable''')
        records = cursor.fetchall()
        p = document.add_paragraph('')
        p.add_run('MEMORY SUMMARY TABLE').bold = True

        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Device Name'
        hdr_cells[1].text = 'Model'
        hdr_cells[2].text = 'Processor Memory Utilization'
        hdr_cells[3].text = 'Top Process'
        hdr_cells[4].text = 'Status'

        for row in records:

            row_cells = table.add_row().cells
            row_cells[0].text = (row[0])
            row_cells[1].text = (row[1])
            row_cells[2].text = (row[2]+"%")
            row_cells[3].text = (row[3])
            row_cells[4].text = (row[4])
        
        #SHOW ENVIRONMENT
        #sql query
        cursor.execute('''SELECT * FROM envtable''')
        records = cursor.fetchall()
        #add to document
        p = document.add_paragraph('')
        p.add_run('Hardware Condition Analysis').bold = True
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Device Name'
        hdr_cells[1].text = 'System'
        hdr_cells[2].text = 'Item'   
        hdr_cells[3].text = 'Status'
        iteration_check = False
        row_check = 'ludesdeveloper'
        count_row = 1       
        for row in records:
            row_cells = table.add_row().cells

            if row_check not in row[1]:
                row_check = row[1]
                iteration_check = False
                start_row = table.cell(count_row, 0)
            else:
                pass

            if row_check == row[1] and iteration_check == False:
                row_cells[0].text = (row[1])
                iteration_check = True
            else:
                row_cells[0].text = ('')
                end_row = table.cell(count_row, 0)
                merge_row = start_row.merge(end_row)

            row_cells[1].text = (row[2]) 
            row_cells[2].text = (row[3])
            row_cells[3].text = (row[4])
            count_row+=1

        #LOG
        #sql query
        cursor.execute('''SELECT * FROM logtable''')
        records = cursor.fetchall()
        #add to document
        p = document.add_paragraph('')
        p.add_run('LOG TABLE CHECKING').bold = True
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Device Name'
        hdr_cells[1].text = 'Model'
        hdr_cells[2].text = 'Script'
        for row in records:
            row_cells = table.add_row().cells
            row_cells[0].text = (row[1])
            row_cells[1].text = (row[2])
            row_cells[2].text = (row[3])
        #close database
        db.close()
        
        #save document
        print('')
        print('Saving Document')
        document.save('preventive_maintenance.docx')
        print('Document has been saved to preventive_maintenance.docx')